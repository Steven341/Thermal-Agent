from pathlib import Path
from typing import Any, Dict
from docx import Document
from docx.shared import Inches
from tools.io_utils import case_path, load_json
from tools.case_state import update_state
from tools.decision_logger import log_decision
from tools.knowledge_db import add_case_to_knowledge_db


def generate_report(project_root: Path, case_id: str) -> Dict[str, Any]:
    cp = case_path(project_root, case_id)
    req = load_json(cp / "work" / "requirements.json")
    result = load_json(cp / "results" / "result_summary.json")
    evaluation = load_json(cp / "work" / "latest_evaluation.json")
    config = load_json(cp / "work" / "simulation_config.json")

    path = cp / "report" / "report.docx"
    doc = Document()
    doc.add_heading(f"热仿真 Agent 报告 - {case_id}", 0)
    doc.add_paragraph("当前默认 mock Ansys，真实交付需替换 solver worker 并由工程师审批。")

    doc.add_heading("1. 需求", level=1)
    for k, v in req.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("2. 最终配置", level=1)
    doc.add_paragraph(f"Solver: {config.get('solver')}")
    doc.add_paragraph(f"Design variables: {config.get('design_variables')}")
    doc.add_paragraph(f"Boundary conditions: {config.get('boundary_conditions')}")

    doc.add_heading("3. 结果", level=1)
    doc.add_paragraph(f"Evaluation: {evaluation.get('evaluation')}")
    doc.add_paragraph(f"Reason: {evaluation.get('reason')}")
    doc.add_paragraph(f"Max temperature: {result.get('max_temperature_c')} °C")
    doc.add_paragraph(f"Target: {result.get('target_max_temperature_c')} °C")
    doc.add_paragraph(f"Passed: {result.get('passed')}")

    doc.add_heading("4. 迭代历史", level=1)
    for d in sorted((cp / "iterations").glob("iter_*")):
        rp = d / "result_summary.json"
        ep = d / "evaluation.json"
        if rp.exists():
            r = load_json(rp)
            doc.add_paragraph(f"{d.name}: max_temp={r.get('max_temperature_c')} °C, passed={r.get('passed')}")
        if ep.exists():
            e = load_json(ep)
            doc.add_paragraph(f"  evaluation={e.get('evaluation')}")

    doc.add_heading("5. 图片", level=1)
    for img in result.get("images", []):
        p = Path(img)
        if p.exists():
            doc.add_paragraph(p.name)
            doc.add_picture(str(p), width=Inches(5.5))

    doc.add_heading("6. 人工复核 Checklist", level=1)
    for item in [
        "材料牌号与物性来源是否正确",
        "热源、入口、出口、流体域 Named Selections 是否正确",
        "几何清理是否误删关键结构",
        "网格质量和残差是否满足交付标准",
        "优化建议是否符合空间、成本、制造约束"
    ]:
        doc.add_paragraph(item)

    doc.save(path)
    update_state(project_root, case_id, status="completed", current_stage="report_generated", last_tool="generate_report", next_tool=None)
    log_decision(project_root, case_id, "python", "report_generated", {"report_path": str(path)})
    add_case_to_knowledge_db(project_root, case_id)
    return {"status": "success", "report_path": str(path)}
